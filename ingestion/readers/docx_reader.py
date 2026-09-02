"""
שכבת הקריאה הראשונית של תהליך קליטת הנתונים
הקובץ קורא מסמכי מקור ומחלץ מהם פסקאות וטבלאות
בשלב הזה הנתונים נשארים גולמיים ללא פירוש או תקנון

הרצה מתיקיית הפרויקט
python -m ingestion.readers.docx_reader
"""

from __future__ import annotations

import sys
from pathlib import Path
from typing import Any

from docx import Document

# מגדיר את תיקיית השורש של הפרויקט
PROJECT_ROOT = Path(__file__).resolve().parents[2]

# מגדיר את התיקייה שבה נמצאים מסמכי המקור
SOURCE_DOCS_DIR = (
    PROJECT_ROOT
    / "data"
    / "raw"
    / "lecturer_samples"
)


def read_docx(file_path: Path) -> dict[str, Any]:
    """
    מקבלת נתיב לקובץ ופותחת אותו
    קוראת את הפסקאות ואת הטבלאות
    מחזירה את התוכן הגולמי ואת שם קובץ המקור
    בשלב הזה עדיין לא מתבצע פירוש של משמעות הנתונים
    """

    document = Document(file_path)

    paragraphs: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            paragraphs.append(text)

    tables: list[list[list[str]]] = []

    for table in document.tables:
        table_rows: list[list[str]] = []

        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
            ]

            table_rows.append(cells)

        tables.append(table_rows)

    return {
        "source_file": file_path.name,
        "paragraphs": paragraphs,
        "tables": tables,
    }


def find_docx_files() -> list[Path]:
    """
    בודקת שתיקיית קובצי המקור קיימת
    מאתרת את כל הקבצים שבתיקייה
    מחזירה רשימה מסודרת של הנתיבים שלהם
    """

    if not SOURCE_DOCS_DIR.exists():
        raise FileNotFoundError(
            f"לא נמצאה תיקיית הנתונים: {SOURCE_DOCS_DIR}"
        )

    return sorted(
        SOURCE_DOCS_DIR.glob("*.docx")
    )


def main() -> None:
    """
    מריצה בדיקה עצמאית של שכבת הקריאה
    מאתרת את הקבצים וקוראת כל אחד מהם
    מדפיסה מידע בסיסי ודוגמאות מהתוכן שנקרא
    משמשת לבדיקה בלבד ולא לניהול תהליך הקליטה המלא
    """

    if sys.platform == "win32":
        try:
            sys.stdout.reconfigure(encoding="utf-8")
        except (AttributeError, OSError):
            pass

    files = find_docx_files()

    print("\n=== DOCX Reader ===")
    print(f"נמצאו {len(files)} קבצי Word.\n")

    for file_path in files:
        data = read_docx(file_path)

        print("—" * 60)
        print("קובץ:", data["source_file"])
        print("מספר פסקאות:", len(data["paragraphs"]))
        print("מספר טבלאות:", len(data["tables"]))

        print("\nדוגמת טקסט:")

        for paragraph in data["paragraphs"][:5]:
            print("-", paragraph)

        if data["tables"]:
            print("\nדוגמה מהטבלה הראשונה:")

            first_table = data["tables"][0]

            for row in first_table[:3]:
                print(row)

        print()


if __name__ == "__main__":
    main()

